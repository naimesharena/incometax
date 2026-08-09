"""Render the detailed computation report as a PDF or Microsoft Word file.

Both renderers consume the structure produced by
:func:`itr1.report.build_computation`.
"""

from __future__ import annotations

import io

ACCENT = (6, 3, 141)        # navy
SAFFRON = (255, 153, 51)
GREEN = (19, 136, 8)


def _fmt(n):
    if n is None:
        return ""
    if isinstance(n, str):
        return n
    sign = "-" if n < 0 else ""
    return f"{sign}{abs(int(round(n))):,}"


def _fmt_rs(n):
    return "" if n is None else f"Rs. {_fmt(n)}" if n else "Rs. 0"


# ---------------------------------------------------------------------------
# Shared section assembly (keeps the PDF / DOCX narratives identical)
# ---------------------------------------------------------------------------

def _part_a_rows(r: dict) -> list:
    """Salary → house property → business → other sources → capital gains →
    gross-total block, honouring the multi-form gti rows when present."""
    rows = [(x["label"], x["note"], x["amount"]) for x in r["salary"]["rows"]]
    pos = min(4, len(rows))
    for ex in r["salary"]["exempt_rows"]:
        rows.insert(pos, (ex["label"], "exempt allowance u/s 10", ex["amount"]))
    rows.append(("Income chargeable under the head 'Salaries'", "", r["salary"]["income"]))
    for hp in r["house_properties"]:
        rows.append((hp["title"], "", None))
        rows.extend((x["label"], x["note"], x["amount"]) for x in hp["rows"])
        rows.append((f"Income from {hp['title'].lower()}", "", hp["income"]))
    if r["house_properties"]:
        rows.append(("Income from house property (total)", r["hp_note"], r["hp_income"]))
    biz = r.get("business")
    if biz:
        rows.append(("Income from business / profession (presumptive scheme)", "", None))
        rows.extend((x["label"], x["note"], x["amount"]) for x in biz["rows"])
    for x in r["other_sources"]["rows"]:
        rows.append((x["label"], x["note"], x["amount"]))
    rows.append(("Income from other sources", "", r["other_sources"]["income"]))
    cg = r.get("capital_gains")
    if cg:
        rows.append(("Income from capital gains (annexure below)", cg["note"],
                     cg["total"]))
    gti = r.get("gti_rows")
    if gti:
        rows.extend((x["label"], x["note"], x["amount"]) for x in gti)
    else:
        rows.append(("Gross total income (before LTCG 112A)", "", r["gti_excl_ltcg"]))
        rows.append(("LTCG u/s 112A (exempt up to Rs. 1,25,000)", r["ltcg_note"],
                     r["ltcg"]))
        rows.append(("Gross total income (incl. LTCG 112A)", "", r["gti_incl_ltcg"]))
    return rows


def _annexure_groups(r: dict) -> list:
    """Optional workings: salary break-up, terminal-benefit least-of-three
    tables, HRA, capital-gain transactions / accrual split, business
    financials, and the Form-10E relief working."""
    groups = []
    sw = r.get("salary_workings")
    if sw:
        rows = []
        for c in sw["components"]:
            note = (f"exempt u/s {c['section']}: Rs. {_fmt(c.get('exempt'))}"
                    if c.get("exempt") else "")
            rows.append((c["label"] or c.get("code", ""), note,
                         (c["amount"] or 0) - (c.get("exempt") or 0)))
        for p in sw["perquisites"]:
            rows.append((f"Perquisite — {p['label']}", "taxable u/s 17(2)", p["amount"]))
        for p in sw["profits"]:
            rows.append((f"Profits in lieu — {p['label']}", "taxable u/s 17(3)",
                         p["amount"]))
        if sw["arrears_total"]:
            rows.append(("of which salary arrears (relief u/s 89 claimed via Form 10E)",
                         "", sw["arrears_total"]))
        if rows:
            groups.append(("Annexure 1 — Salary break-up as per the salary slip",
                           ("Component", "Working / basis", "Taxable (Rs.)"), rows))
        if sw["hra_rows"]:
            groups.append(("Annexure 1A — HRA exemption u/s 10(13A) (least of three)",
                           ("Particulars", "", "Amount (Rs.)"),
                           [(x["label"], "", x["amount"]) for x in sw["hra_rows"]]))
        for tw in sw["terminal"]:
            trows = [(lbl, "", val) for lbl, val in tw.get("lines", [])]
            trows.append(("Received", "", tw.get("received")))
            trows.append(("Less: exempt", tw.get("limit", ""), -(tw.get("exempt") or 0)))
            trows.append(("Taxable (added to salary)", "", tw.get("taxable")))
            groups.append((f"Annexure 1B — Terminal benefit: {tw.get('name')}",
                           ("Particulars", "Working", "Amount (Rs.)"), trows))
    cg = r.get("capital_gains")
    if cg:
        wrows = [(x["label"], "", x["amount"]) for x in cg["rows"]]
        if wrows:
            groups.append(("Annexure 2 — Capital gains computation (head-wise)",
                           ("Particulars", "", "Amount (Rs.)"), wrows))
        if cg.get("details"):
            groups.append(("Annexure 2A — Transaction-wise gains",
                           ("Transaction", "Working / basis", "Gain (Rs.)"),
                           [(x["label"], x["note"], x["amount"]) for x in cg["details"]]))
        if cg.get("quarterly"):
            groups.append(("Annexure 2B — Accrual across the year (for sec 234C)",
                           ("Period", "", "Gain (Rs.)"),
                           [(x["label"], "", x["amount"]) for x in cg["quarterly"]]))
    biz = r.get("business")
    if biz and biz.get("financials_rows"):
        groups.append(("Annexure 3 — No-books case: financial particulars (sec 44AA(2))",
                       ("Particulars", "", "Amount (Rs.)"),
                       [(x["label"], "", x["amount"]) for x in biz["financials_rows"]]))
    if r.get("form10e_rows"):
        meta = r.get("form10e_meta") or {}
        frows = []
        for x in r["form10e_rows"]:
            frows.append((f"Arrear salary of FY {x.get('fy')}",
                          f"arrear {_fmt_rs(x.get('amount'))} · income that year "
                          f"{_fmt_rs(x.get('income_that_fy'))}",
                          x.get("delta")))
        frows.append(("Extra tax of the current year due to arrears", "",
                      meta.get("current_delta")))
        frows.append(("Total extra tax of the earlier years", "",
                      meta.get("sum_prior_deltas")))
        frows.append(("Relief u/s 89(1) — Form 10E", "", meta.get("relief89")))
        groups.append(("Annexure 4 — Relief u/s 89 on salary arrears (Form 10E)",
                       ("Particulars", "Working", "Extra tax (Rs.)"), frows))
    return groups


# ---------------------------------------------------------------------------
# PDF (ReportLab Platypus)
# ---------------------------------------------------------------------------

def computation_pdf(comp: dict) -> bytes:
    """Render ``comp`` as a PDF, returned as raw bytes."""
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
    from reportlab.lib.units import mm
    from reportlab.platypus import (Paragraph, SimpleDocTemplate, Spacer,
                                    Table, TableStyle)

    buf = io.BytesIO()
    doc = SimpleDocTemplate(
        buf, pagesize=A4, leftMargin=16 * mm, rightMargin=16 * mm,
        topMargin=14 * mm, bottomMargin=14 * mm,
        title=f"ITR-1 Computation AY {comp['meta']['assessment_year']}",
        author=comp["meta"]["software"])

    styles = getSampleStyleSheet()
    s_title = ParagraphStyle("t", parent=styles["Title"], fontSize=16,
                             textColor=colors.Color(*(c / 255 for c in ACCENT)),
                             spaceAfter=2 * mm, alignment=1)
    s_sub = ParagraphStyle("s", parent=styles["Normal"], fontSize=9,
                           textColor=colors.HexColor("#667085"), alignment=1,
                           spaceAfter=4 * mm)
    s_h1 = ParagraphStyle("h1", parent=styles["Heading2"], fontSize=11.5,
                          textColor=colors.Color(*(c / 255 for c in ACCENT)),
                          spaceBefore=4 * mm, spaceAfter=1.5 * mm)
    s_h2 = ParagraphStyle("h2", parent=styles["Heading3"], fontSize=10,
                          textColor=colors.HexColor("#1c2434"),
                          spaceBefore=2.5 * mm, spaceAfter=1 * mm)
    s_note = ParagraphStyle("n", parent=styles["Normal"], fontSize=8,
                            textColor=colors.HexColor("#667085"))

    def para(text, style=s_note):
        return Paragraph(str(text), style)

    def money_table(rows, col_labels=("Particulars", "Working / basis", "Amount (Rs.)"),
                    widths=(95 * mm, 55 * mm, 28 * mm)):
        data = [[para(f"<b>{c}</b>", s_note) for c in col_labels]]
        for r in rows:
            lbl, note, amt = r[0], (r[1] if len(r) > 1 else ""), (r[2] if len(r) > 2 else None)
            amt_txt = "" if amt is None else _fmt(amt)
            data.append([para(lbl), para(note), para(f"<b>{amt_txt}</b>" if amt_txt else "")])
        t = Table(data, colWidths=widths, repeatRows=1)
        t.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (-1, 0), colors.Color(*(c / 255 for c in ACCENT)),
             ),
            ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
            ("GRID", (0, 0), (-1, -1), 0.25, colors.HexColor("#dde1e8")),
            ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#f7f8fc")]),
            ("ALIGN", (2, 0), (2, -1), "RIGHT"),
            ("VALIGN", (0, 0), (-1, -1), "TOP"),
            ("LEFTPADDING", (0, 0), (-1, -1), 4),
            ("RIGHTPADDING", (0, 0), (-1, -1), 4),
            ("TOPPADDING", (0, 0), (-1, -1), 2.5),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 2.5),
        ]))
        return t

    story = []
    meta = comp["meta"]
    story.append(para("COMPUTATION OF TOTAL INCOME AND TAX LIABILITY", s_title))
    story.append(para(
        f"{meta['form']} &nbsp;·&nbsp; Assessment Year {meta['assessment_year']} "
        f"(Previous Year {meta['previous_year']})", s_sub))
    story.append(money_table([
        ("Taxpayer", "", meta["name"] or "-"),
        ("PAN", "", meta["pan"] or "-"),
        ("Return filed u/s", "", meta["return_section"]),
        ("Generated on", "", meta["generated_on"]),
        ("Prepared with", "", meta["software"]),
    ], col_labels=("Return particulars", "", "")))
    story.append(Spacer(1, 3 * mm))

    # ---- comparison ----
    story.append(para("Regime comparison", s_h1))
    cp = comp["comparison"]
    rec = "New" if cp["recommended"] == "new" else "Old"
    story.append(money_table([
        ("New tax regime (default u/s 115BAC) - total tax + interest", "",
         cp["new_total"],),
        ("Old tax regime (with deductions) - total tax + interest", "", cp["old_total"]),
        ("Recommendation", f"{rec} regime is beneficial", cp["saving"]),
    ]))
    story.append(para(
        "Amount in the recommendation row is the saving offered by the beneficial regime.",
        s_note))

    for regime_key in ("new", "old"):
        r = comp[regime_key]
        story.append(para(r["regime_title"].upper(), s_h1))
        if r["age_note"]:
            story.append(para(r["age_note"], s_note))

        # --- Part A ---
        story.append(para("PART A - Computation of Total Income", s_h2))
        story.append(money_table(_part_a_rows(r)))

        if r["via_rows"]:
            story.append(para("Deductions under Chapter VI-A", s_h2))
            if r["via_caption"]:
                story.append(para(r["via_caption"], s_note))
            vrows = [(f"<b>{x['section']}</b>", x["note"], x["eligible"]) for x in r["via_rows"]]
            vrows.append(("<b>Total deductions under Chapter VI-A</b>",
                          "restricted to gross total income", r["via_total"]))
            story.append(money_table(vrows, col_labels=("Section", "Details", "Allowable (Rs.)")))

        story.append(money_table([
            ("<b>TOTAL INCOME</b>", r["total_income_note"], r["total_income"]),
        ]))

        for a_title, a_labels, a_rows in _annexure_groups(r):
            story.append(para(a_title, s_h2))
            story.append(money_table(a_rows, col_labels=a_labels))

        # --- Part B ---
        story.append(para("PART B - Computation of Tax Liability", s_h2))
        if r["slab_rows"]:
            srows = [(x["slab"], f"{x['rate'] * 100:g}% on Rs. {_fmt(x['taxable'])}",
                      x["tax"]) for x in r["slab_rows"]]
            story.append(money_table(srows, col_labels=("Slab (Rs.)", "Rate × taxable",
                                                        "Tax (Rs.)"),
                                     widths=(60 * mm, 90 * mm, 28 * mm)))
        brow = [
            ("Tax at slab rates (on normal-rate income)", "", r["slab_tax"]),
        ]
        if r.get("special_tax"):
            brow.append(("Tax at special rates (111A @ 20% / LTCG @ 12.5% or 20%)",
                         r.get("special_tax_note", ""), r["special_tax"]))
        brow += [
            ("Tax on LTCG u/s 112A (above Rs. 1,25,000)", "", r["ltcg_tax"]),
            ("Surcharge", r["surcharge_note"], r["surcharge"]),
            ("<b>Tax payable on total income</b>", "", r["tax_payable"]),
            ("Less: rebate u/s 87A", r["rebate_note"], -r["rebate"]),
            ("Tax after rebate", "", r["tax_after_rebate"]),
            ("Add: Health & Education Cess @ 4%", r["cess_note"], r["cess"]),
            ("<b>Gross tax liability</b>", "", r["gross_tax_liability"]),
            ("Less: relief u/s 89", r["relief89_note"], -r["relief89"]),
            ("Net tax liability (rounded u/s 288B)", "", r["net_tax_liability"]),
        ]
        for it in r["interest"]:
            brow.append((it["label"], it["note"], it["amount"]))
            for inst in it.get("instalments", []) or []:
                brow.append((
                    f"    instalment due {inst['due']} ({inst['required_pct'] * 100:g}%)",
                    f"required Rs. {_fmt(inst['required'])}, paid Rs. {_fmt(inst['paid'])}, "
                    f"1% × {inst['months']} mo.",
                    inst["amount"]))
        brow.append(("<b>Total tax + interest & fee</b>", "", r["tot_tax_plus_interest"]))
        story.append(money_table(brow))

        # --- Part C ---
        story.append(para("PART C - Taxes Paid and Refund / Balance Payable", s_h2))
        tp = r["taxes_paid"]
        crows = [
            ("Tax deducted at source (TDS-1 + TDS-2 + TDS-3)", "", tp["tds"]),
            ("Tax collected at source (TCS)", "", tp["tcs"]),
            ("Advance tax", "challans deposited within the previous year", tp["advance_tax"]),
            ("Self-assessment tax", "challans deposited after 31-03-2026", tp["self_assessment_tax"]),
            ("<b>Total taxes paid</b>", "", tp["total"]),
            ("<b>Balance tax payable</b>", "", r["balance_payable"]),
            ("<b>Refund due</b>", "", r["refund_due"]),
        ]
        story.append(money_table(crows))
        if r["result_note"]:
            story.append(para(r["result_note"], s_note))

        if r["warnings"]:
            story.append(para("Notes &amp; adjustments applied automatically", s_h2))
            for w in r["warnings"]:
                story.append(para(f"• {w}", s_note))

    story.append(Spacer(1, 4 * mm))
    story.append(para(
        "This computation sheet is generated for assistance and should be verified "
        "against Form 16/16A, Form 26AS and AIS before filing the return.", s_note))

    doc.build(story)
    return buf.getvalue()


# ---------------------------------------------------------------------------
# DOCX (python-docx)
# ---------------------------------------------------------------------------

def computation_docx(comp: dict) -> bytes:
    """Render ``comp`` as a .docx (Microsoft Word) file, returned as bytes."""
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    document = Document()
    style = document.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10)

    meta = comp["meta"]
    title = document.add_heading("Computation of Total Income and Tax Liability", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = document.add_paragraph(
        f"{meta['form']}  ·  Assessment Year {meta['assessment_year']} "
        f"(Previous Year {meta['previous_year']})")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in p.runs:
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x66, 0x70, 0x85)

    def kv_table(rows, headers=("Particulars", "Working / basis", "Amount (Rs.)")):
        table = document.add_table(rows=1, cols=len(headers))
        table.style = "Light Grid Accent 1"
        for i, h in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.text = h
            for par in cell.paragraphs:
                for run in par.runs:
                    run.bold = True
        for row in rows:
            lbl, note, amt = row[0], (row[1] if len(row) > 1 else ""), (row[2] if len(row) > 2 else None)
            cells = table.add_row().cells
            cells[0].text = str(lbl)
            cells[1].text = str(note or "")
            cells[2].text = "" if amt is None else (amt if isinstance(amt, str) else f"{int(round(amt)):,}")
        for row in table.rows[1:]:
            row.cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        return table

    info = [("Taxpayer", "", meta["name"] or "-"), ("PAN", "", meta["pan"] or "-"),
            ("Return filed u/s", "", meta["return_section"]),
            ("Generated on", "", meta["generated_on"]),
            ("Prepared with", "", meta["software"])]
    kv_table(info)

    document.add_heading("Regime comparison", 1)
    cp_ = comp["comparison"]
    rec = "New" if cp_["recommended"] == "new" else "Old"
    kv_table([
        ("New tax regime (default u/s 115BAC) - total tax + interest", "", cp_["new_total"]),
        ("Old tax regime (with deductions) - total tax + interest", "", cp_["old_total"]),
        ("Recommendation", f"{rec} regime is beneficial", cp_["saving"]),
    ])

    for regime_key in ("new", "old"):
        r = comp[regime_key]
        document.add_heading(r["regime_title"], 1)
        if r["age_note"]:
            document.add_paragraph(r["age_note"]).runs[0].italic = True

        document.add_heading("PART A - Computation of Total Income", 2)
        kv_table(_part_a_rows(r))

        if r["via_rows"]:
            document.add_heading("Deductions under Chapter VI-A", 2)
            if r["via_caption"]:
                document.add_paragraph(r["via_caption"]).runs[0].italic = True
            vrows = [(x["section"], x["note"], x["eligible"]) for x in r["via_rows"]]
            vrows.append(("Total deductions under Chapter VI-A",
                          "restricted to gross total income", r["via_total"]))
            kv_table(vrows, headers=("Section", "Details", "Allowable (Rs.)"))

        kv_table([("TOTAL INCOME", r["total_income_note"], r["total_income"])])

        for a_title, a_labels, a_rows in _annexure_groups(r):
            document.add_heading(a_title, 2)
            kv_table(a_rows, headers=a_labels)

        document.add_heading("PART B - Computation of Tax Liability", 2)
        if r["slab_rows"]:
            srows = [(x["slab"], f"{x['rate'] * 100:g}% on Rs. {int(round(x['taxable'])):,}",
                      x["tax"]) for x in r["slab_rows"]]
            kv_table(srows, headers=("Slab (Rs.)", "Rate × taxable", "Tax (Rs.)"))
        brow = [
            ("Tax at slab rates (on normal-rate income)", "", r["slab_tax"]),
        ]
        if r.get("special_tax"):
            brow.append(("Tax at special rates (111A @ 20% / LTCG @ 12.5% or 20%)",
                         r.get("special_tax_note", ""), r["special_tax"]))
        brow += [
            ("Tax on LTCG u/s 112A (above Rs. 1,25,000)", "", r["ltcg_tax"]),
            ("Surcharge", r["surcharge_note"], r["surcharge"]),
            ("Tax payable on total income", "", r["tax_payable"]),
            ("Less: rebate u/s 87A", r["rebate_note"], -r["rebate"]),
            ("Tax after rebate", "", r["tax_after_rebate"]),
            ("Add: Health & Education Cess @ 4%", r["cess_note"], r["cess"]),
            ("Gross tax liability", "", r["gross_tax_liability"]),
            ("Less: relief u/s 89", r["relief89_note"], -r["relief89"]),
            ("Net tax liability (rounded u/s 288B)", "", r["net_tax_liability"]),
        ]
        for it in r["interest"]:
            brow.append((it["label"], it["note"], it["amount"]))
            for inst in it.get("instalments", []) or []:
                brow.append((
                    f"    instalment due {inst['due']} ({inst['required_pct'] * 100:g}%)",
                    f"required Rs. {int(round(inst['required'])):,}, paid Rs. "
                    f"{int(round(inst['paid'])):,}, 1% × {inst['months']} mo.",
                    inst["amount"]))
        brow.append(("Total tax + interest & fee", "", r["tot_tax_plus_interest"]))
        kv_table(brow)

        document.add_heading("PART C - Taxes Paid and Refund / Balance Payable", 2)
        tp = r["taxes_paid"]
        kv_table([
            ("Tax deducted at source (TDS-1 + TDS-2 + TDS-3)", "", tp["tds"]),
            ("Tax collected at source (TCS)", "", tp["tcs"]),
            ("Advance tax", "challans deposited within the previous year", tp["advance_tax"]),
            ("Self-assessment tax", "challans deposited after 31-03-2026", tp["self_assessment_tax"]),
            ("Total taxes paid", "", tp["total"]),
            ("Balance tax payable", "", r["balance_payable"]),
            ("Refund due", "", r["refund_due"]),
        ])
        if r["result_note"]:
            document.add_paragraph(r["result_note"]).runs[0].italic = True

        if r["warnings"]:
            document.add_heading("Notes & adjustments applied automatically", 2)
            for w in r["warnings"]:
                document.add_paragraph(w, style="List Bullet")

    document.add_paragraph(
        "This computation sheet is generated for assistance and should be verified "
        "against Form 16/16A, Form 26AS and AIS before filing the return.")

    buf = io.BytesIO()
    document.save(buf)
    return buf.getvalue()
